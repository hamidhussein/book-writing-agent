from __future__ import annotations

import base64
import io
import re
from typing import Any, Dict, List

from django.utils import timezone

from ..models import BookProject, Chapter, ChapterStatus, ProjectStatus
from .llm import LLMService
from .vector_store import VectorMemoryStore


class BookWorkflowService:
    def __init__(self) -> None:
        self.llm = LLMService()
        self.vector_store = VectorMemoryStore(self.llm)

    def execute_mode(self, project: BookProject, mode: str, inputs: Dict[str, Any]) -> Dict[str, Any]:
        if mode == "toc":
            return self._run_toc(project)
        if mode == "refine_toc":
            return self._run_refine_toc(project, inputs)
        if mode == "chapter":
            return self._run_chapter(project, inputs)
        if mode == "export":
            return self._run_export(project, inputs)
        if mode == "profile_assistant":
            return self._run_profile_assistant(project, inputs)
        raise ValueError("mode must be one of: toc | refine_toc | chapter | export | profile_assistant")

    def _run_toc(self, project: BookProject) -> Dict[str, Any]:
        kb_context = self.vector_store.search_knowledge_base(
            project_id=str(project.id),
            query=f"{project.title} {project.genre} {project.target_audience} book plan",
            limit=8,
        )
        payload = self.llm.generate_outline(project, knowledge_context=kb_context)
        outline = self._normalize_outline(payload.get("outline", {}))
        fallback_info = self._runtime_fallback_info(payload)
        outline_profile_compliance = self._outline_profile_compliance(project, outline)
        llm_runtime_meta = self._merge_dicts(
            payload.get("metadata", {}),
            {"profile_compliance": outline_profile_compliance},
        )

        project.outline_json = outline
        project.metadata_json = self._merge_project_metadata(project, llm_runtime_meta)
        project.status = ProjectStatus.OUTLINED
        project.updated_at = timezone.now()
        project.save(update_fields=["outline_json", "metadata_json", "status", "updated_at"])

        self._sync_chapters_from_outline(project, outline)

        warnings = list(outline_profile_compliance.get("issues", [])) if outline_profile_compliance.get("fail") else []

        response = {
            "status": "success",
            "outline": outline,
            "metadata": project.metadata_json,
            "used_fallback": fallback_info["used_fallback"],
            "fallback_stages": fallback_info["fallback_stages"],
            "next_steps": payload.get(
                "next_steps",
                [
                    "Review the generated outline.",
                    "Refine the outline with targeted feedback.",
                    "Generate chapters one-by-one once approved.",
                ],
            ),
        }
        if warnings:
            response["warnings"] = warnings
        return response

    def _run_refine_toc(self, project: BookProject, inputs: Dict[str, Any]) -> Dict[str, Any]:
        feedback = str(inputs.get("feedback", "")).strip()
        if not feedback:
            raise ValueError("feedback is required for refine_toc mode")
        refine_feedback_analysis = self._analyze_refine_feedback_conflicts(project, feedback)
        existing_outline = project.outline_json or {}
        kb_context = self.vector_store.search_knowledge_base(
            project_id=str(project.id),
            query=f"{project.title} outline refinement {feedback}",
            limit=8,
        )
        payload = self.llm.refine_outline(
            project,
            existing_outline,
            feedback,
            knowledge_context=kb_context,
        )
        outline = self._normalize_outline(payload.get("outline", {}))
        fallback_info = self._runtime_fallback_info(payload)
        outline_profile_compliance = self._outline_profile_compliance(project, outline)
        llm_runtime_meta = self._merge_dicts(
            payload.get("metadata", {}),
            {
                "profile_compliance": outline_profile_compliance,
                "refine_feedback_analysis": refine_feedback_analysis,
            },
        )

        project.outline_json = outline
        project.metadata_json = self._merge_project_metadata(project, llm_runtime_meta)
        project.status = ProjectStatus.OUTLINED
        project.updated_at = timezone.now()
        project.save(update_fields=["outline_json", "metadata_json", "status", "updated_at"])

        self._sync_chapters_from_outline(project, outline)

        warnings: List[str] = []
        if bool(refine_feedback_analysis.get("warn")):
            warnings.extend([str(issue) for issue in refine_feedback_analysis.get("issues", []) if str(issue).strip()])
        if outline_profile_compliance.get("fail"):
            warnings.extend([str(issue) for issue in outline_profile_compliance.get("issues", []) if str(issue).strip()])
        deduped_warnings: List[str] = []
        for warning in warnings:
            if warning not in deduped_warnings:
                deduped_warnings.append(warning)

        response = {
            "status": "success",
            "outline": outline,
            "metadata": project.metadata_json,
            "used_fallback": fallback_info["used_fallback"],
            "fallback_stages": fallback_info["fallback_stages"],
            "next_steps": payload.get(
                "next_steps",
                [
                    "Review the refined outline.",
                    "Generate the next chapter.",
                ],
            ),
        }
        if deduped_warnings:
            response["warnings"] = deduped_warnings
        return response

    def _run_profile_assistant(self, project: BookProject, inputs: Dict[str, Any]) -> Dict[str, Any]:
        message = str(inputs.get("message", "")).strip()
        current_profile = inputs.get("current_profile", {})
        if not isinstance(current_profile, dict):
            current_profile = {}
        conversation = inputs.get("conversation", [])
        if not isinstance(conversation, list):
            conversation = []

        payload = self.llm.assist_profile(
            project=project,
            current_profile=current_profile,
            conversation=conversation,
            user_message=message,
        )
        self._apply_profile_updates_if_finalized(project, payload)
        
        fallback_info = self._runtime_fallback_info(payload)
        return {
            "status": "success",
            "assistant_response": payload,
            "used_fallback": fallback_info["used_fallback"],
            "fallback_stages": fallback_info["fallback_stages"],
        }

    def _apply_profile_updates_if_finalized(self, project: BookProject, payload: Dict[str, Any]) -> None:
        if not isinstance(payload, dict):
            return
        if not bool(payload.get("is_finalized")):
            return

        updates = payload.get("field_updates", {})
        if not isinstance(updates, dict) or not updates:
            return

        project_updates: Dict[str, Any] = {}
        if isinstance(updates.get("title"), str) and updates["title"].strip():
            project_updates["title"] = updates["title"].strip()[:160]
        if isinstance(updates.get("genre"), str) and updates["genre"].strip():
            project_updates["genre"] = updates["genre"].strip()[:80]
        if isinstance(updates.get("audience"), str) and updates["audience"].strip():
            project_updates["target_audience"] = updates["audience"].strip()[:80]
        if isinstance(updates.get("language"), str) and updates["language"].strip():
            project_updates["language"] = updates["language"].strip()[:40]
        if isinstance(updates.get("tone"), str) and updates["tone"].strip():
            project_updates["tone"] = updates["tone"].strip()[:80]
        if "length" in updates:
            try:
                project_updates["target_word_count"] = max(300, int(float(str(updates["length"]).strip())))
            except Exception:
                pass

        raw_meta = project.metadata_json if isinstance(project.metadata_json, dict) else {}
        user_concept = raw_meta.get("user_concept", {})
        if not isinstance(user_concept, dict):
            user_concept = {}
        profile = user_concept.get("profile", {})
        if not isinstance(profile, dict):
            profile = {}
        profile.update({str(k): v for k, v in updates.items()})
        user_concept["profile"] = profile

        new_meta = dict(raw_meta)
        new_meta["user_concept"] = user_concept
        new_meta["profile"] = profile
        project_updates["metadata_json"] = new_meta

        if project_updates:
            for field, value in project_updates.items():
                setattr(project, field, value)
            project.save(update_fields=list(project_updates.keys()) + ["updated_at"])

    def _run_chapter(self, project: BookProject, inputs: Dict[str, Any]) -> Dict[str, Any]:
        chapter_ctx = self.prepare_chapter_context(project, inputs)
        payload = self.llm.generate_chapter(
            project,
            chapter_ctx["outline"],
            chapter_ctx["chapter_number"],
            memory_context=chapter_ctx["memory_context"],
            knowledge_context=chapter_ctx["knowledge_context"],
        )
        fallback_info = self._runtime_fallback_info(payload)
        return self.persist_chapter_result(
            project=project,
            outline=chapter_ctx["outline"],
            chapter_number=chapter_ctx["chapter_number"],
            target=chapter_ctx["target"],
            chapter_data=payload.get("chapter", {}),
            metadata=self._merge_dicts(payload.get("metadata", {}), fallback_info),
            next_steps=payload.get(
                "next_steps",
                ["Review the chapter and proceed to the next one."],
            ),
        )

    def prepare_chapter_context(self, project: BookProject, inputs: Dict[str, Any]) -> Dict[str, Any]:
        outline = self._normalize_outline(project.outline_json or {})
        if not outline.get("chapters"):
            raise ValueError("Project does not have an outline yet")

        chapter_number = self._to_int(inputs.get("chapter_number"), "chapter_number")
        target = next((c for c in outline["chapters"] if c["number"] == chapter_number), None)
        if not target:
            raise ValueError("chapter_number is outside outline range")

        memory_context = self.vector_store.search_memory(
            project_id=str(project.id),
            query=f"{project.title} chapter {chapter_number} continuity and style",
            limit=5,
        )
        knowledge_context = self.vector_store.search_knowledge_base(
            project_id=str(project.id),
            query=f"{project.title} chapter {chapter_number} facts references examples",
            limit=6,
        )

        return {
            "outline": outline,
            "chapter_number": chapter_number,
            "target": target,
            "memory_context": memory_context,
            "knowledge_context": knowledge_context,
        }

    def persist_chapter_result(
        self,
        project: BookProject,
        outline: Dict[str, Any],
        chapter_number: int,
        target: Dict[str, Any],
        chapter_data: Dict[str, Any],
        metadata: Dict[str, Any] | None = None,
        next_steps: List[str] | None = None,
    ) -> Dict[str, Any]:
        content = str(chapter_data.get("content", "")).strip()
        if not content:
            raise ValueError("Generated chapter content is empty")

        chapter, _ = Chapter.objects.update_or_create(
            project=project,
            number=chapter_number,
            defaults={
                "title": str(chapter_data.get("title") or target.get("title", "")).strip() or f"Chapter {chapter_number}",
                "content": content,
                "summary": str(chapter_data.get("summary", "")).strip(),
                "status": ChapterStatus.GENERATED,
            },
        )

        vector_indexed = self.vector_store.upsert_chapter_memory(
            project_id=str(project.id),
            chapter_number=chapter.number,
            title=chapter.title,
            content=chapter.content,
            summary=chapter.summary,
        )
        if vector_indexed and not chapter.vector_indexed:
            chapter.vector_indexed = True
            chapter.save(update_fields=["vector_indexed", "updated_at"])

        project.status = ProjectStatus.WRITING
        project.updated_at = timezone.now()
        project.save(update_fields=["status", "updated_at"])

        return {
            "status": "success",
            "outline": outline,
            "chapter": {
                "number": chapter.number,
                "title": chapter.title,
                "content": chapter.content,
                "summary": chapter.summary,
            },
            "metadata": metadata if isinstance(metadata, dict) else {},
            "next_steps": next_steps if isinstance(next_steps, list) and next_steps else ["Review the chapter and proceed to the next one."],
        }

    def _run_export(self, project: BookProject, inputs: Dict[str, Any]) -> Dict[str, Any]:
        outline = self._normalize_outline(project.outline_json or {})
        if not outline.get("chapters"):
            raise ValueError("Project does not have an outline yet")

        export_format = str(inputs.get("export_format", "pdf")).strip().lower()
        if export_format not in {"pdf", "docx", "both"}:
            raise ValueError("export_format must be one of: pdf | docx | both")

        chapters = list(project.chapters.all().order_by("number"))
        if not chapters:
            raise ValueError("No chapters found. Generate at least one chapter before export.")

        warnings: List[str] = []
        output: Dict[str, Any] = {
            "status": "success",
            "warnings": warnings,
            "used_fallback": False,
            "fallback_stages": [],
        }

        chapter_payload = [
            {"number": c.number, "title": c.title, "content": c.content, "summary": c.summary}
            for c in chapters
        ]
        if any(self._has_visual_placeholders(ch.get("content", "")) for ch in chapter_payload):
            warnings.append(
                "Figure/flowchart placeholders were preserved in export. Render visual assets in a post-processing step to replace placeholders."
            )

        if export_format in {"pdf", "both"}:
            pdf_bytes = self._render_pdf(project, outline, chapter_payload)
            output["pdf_base64"] = base64.b64encode(pdf_bytes).decode("utf-8")
            output["pdf_filename"] = f"{self._safe_file(project.title)}.pdf"

        if export_format in {"docx", "both"}:
            docx_bytes = self._render_docx(project, outline, chapter_payload)
            output["docx_base64"] = base64.b64encode(docx_bytes).decode("utf-8")
            output["docx_filename"] = f"{self._safe_file(project.title)}.docx"

        project.status = ProjectStatus.EXPORTED
        project.updated_at = timezone.now()
        project.save(update_fields=["status", "updated_at"])

        output["outline"] = outline
        output["next_steps"] = [
            "Download exported file(s).",
            "Review formatting in your editor.",
            "Publish or continue editing.",
        ]
        return output

    def _runtime_fallback_info(self, payload: Dict[str, Any] | Any) -> Dict[str, Any]:
        if not isinstance(payload, dict):
            return {"used_fallback": False, "fallback_stages": []}
        used_fallback = bool(payload.get("used_fallback"))
        stage = str(payload.get("fallback_stage", "")).strip()
        stages = [stage] if used_fallback and stage else []
        return {"used_fallback": used_fallback, "fallback_stages": stages}

    def _merge_dicts(self, left: Dict[str, Any] | Any, right: Dict[str, Any] | Any) -> Dict[str, Any]:
        out: Dict[str, Any] = {}
        if isinstance(left, dict):
            out.update(left)
        if isinstance(right, dict):
            out.update(right)
        return out

    def _sync_chapters_from_outline(self, project: BookProject, outline: Dict[str, Any]) -> None:
        for chapter in outline.get("chapters", []):
            Chapter.objects.get_or_create(
                project=project,
                number=int(chapter["number"]),
                defaults={
                    "title": str(chapter["title"]).strip(),
                    "content": "",
                    "summary": "",
                    "status": ChapterStatus.PENDING,
                },
            )

    def _normalize_outline(self, outline: Dict[str, Any]) -> Dict[str, Any]:
        if not isinstance(outline, dict):
            raise ValueError("outline must be an object")

        chapters = outline.get("chapters", [])
        if not isinstance(chapters, list) or not chapters:
            raise ValueError("outline.chapters must be a non-empty array")

        normalized = []
        expected = 1
        for chapter in chapters:
            if not isinstance(chapter, dict):
                raise ValueError("outline chapter must be an object")
            number = self._to_int(chapter.get("number"), "outline.chapter.number")
            if number != expected:
                raise ValueError("outline chapter numbers must be sequential starting at 1")
            title = str(chapter.get("title", "")).strip()
            if not title:
                raise ValueError(f"outline chapter {expected} missing title")
            bullet_points = chapter.get("bullet_points", [])
            if not isinstance(bullet_points, list):
                raise ValueError(f"outline chapter {expected} bullet_points must be an array")
            normalized.append(
                {
                    "number": number,
                    "title": title,
                    "bullet_points": [str(p).strip() for p in bullet_points if str(p).strip()],
                }
            )
            expected += 1

        synopsis = str(outline.get("synopsis", "")).strip()
        return {"synopsis": synopsis, "chapters": normalized}

    def _to_int(self, value: Any, field: str) -> int:
        if value is None or str(value).strip() == "":
            raise ValueError(f"{field} is required")
        try:
            return int(float(str(value).strip()))
        except Exception as exc:
            raise ValueError(f"{field} must be a valid integer") from exc

    def _render_pdf(self, project: BookProject, outline: Dict[str, Any], chapters: List[Dict[str, Any]]) -> bytes:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import inch
        from reportlab.platypus import PageBreak, Paragraph, Preformatted, SimpleDocTemplate, Spacer

        buf = io.BytesIO()
        doc = SimpleDocTemplate(
            buf,
            pagesize=A4,
            topMargin=1 * inch,
            bottomMargin=1 * inch,
            leftMargin=1 * inch,
            rightMargin=1 * inch,
        )
        styles = getSampleStyleSheet()
        h1 = ParagraphStyle("BookH1", parent=styles["Heading1"], spaceAfter=12)
        h2 = ParagraphStyle("BookH2", parent=styles["Heading2"], spaceAfter=10)
        body = ParagraphStyle("BookBody", parent=styles["BodyText"], leading=14, spaceAfter=8)
        quote_style = ParagraphStyle("BookQuote", parent=body, leftIndent=18, rightIndent=8, italic=True)
        callout_style = ParagraphStyle("BookCallout", parent=body, leftIndent=18, rightIndent=8, backColor="#f3f4f6")
        placeholder_style = ParagraphStyle("BookPlaceholder", parent=body, italic=True, textColor="#555555")
        code_style = ParagraphStyle(
            "BookCode",
            parent=styles["BodyText"],
            fontName="Courier",
            fontSize=9,
            leading=11,
            leftIndent=12,
            rightIndent=6,
            spaceAfter=8,
        )

        story: List[Any] = []
        story.append(Spacer(1, 2 * inch))
        story.append(Paragraph(self._escape(project.title), styles["Title"]))
        story.append(Spacer(1, 0.2 * inch))
        story.append(Paragraph(self._escape(f"{project.genre}  {project.language}  {project.tone}"), styles["Italic"]))
        story.append(PageBreak())

        story.append(Paragraph("Table of Contents", h1))
        for ch in outline.get("chapters", []):
            story.append(Paragraph(self._escape(f"Chapter {ch['number']}: {ch['title']}"), body))
        story.append(PageBreak())

        for ch in chapters:
            story.append(Paragraph(self._escape(f"Chapter {ch['number']}"), h2))
            story.append(Paragraph(self._escape(ch["title"]), h1))
            story.append(Spacer(1, 0.1 * inch))
            for block in self._iter_render_blocks(ch["content"]):
                block_type = str(block.get("type", "paragraph"))
                block_text = str(block.get("text", "")).strip()
                if not block_text and block_type not in {"visual_placeholder"}:
                    continue
                if block_type == "h1":
                    story.append(Paragraph(self._escape(block_text), h1))
                elif block_type == "h2":
                    story.append(Paragraph(self._escape(block_text), h2))
                elif block_type == "code":
                    story.append(Preformatted(block_text, code_style))
                elif block_type == "quote":
                    story.append(Paragraph(self._escape(block_text), quote_style))
                elif block_type == "callout":
                    label = str(block.get("label", "Note")).strip() or "Note"
                    story.append(Paragraph(self._escape(f"{label}: {block_text}"), callout_style))
                elif block_type == "visual_placeholder":
                    kind = str(block.get("kind", "visual")).upper()
                    label = str(block.get("label", "")).strip()
                    story.append(Paragraph(self._escape(f"[{kind} PLACEHOLDER] {label}"), placeholder_style))
                elif block_type == "table":
                    story.append(Preformatted(block_text, code_style))
                else:
                    story.append(Paragraph(self._escape(block_text), body))
            story.append(PageBreak())

        doc.build(story)
        buf.seek(0)
        return buf.read()

    def _render_docx(self, project: BookProject, outline: Dict[str, Any], chapters: List[Dict[str, Any]]) -> bytes:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt

        document = Document()
        title_paragraph = document.add_paragraph()
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_paragraph.add_run(project.title)
        run.bold = True
        run.font.size = Pt(26)

        sub = document.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_run = sub.add_run(f"{project.genre}  {project.language}  {project.tone}")
        sub_run.italic = True
        document.add_page_break()

        document.add_heading("Table of Contents", level=1)
        for ch in outline.get("chapters", []):
            document.add_paragraph(f"Chapter {ch['number']}: {ch['title']}")
        document.add_page_break()

        for ch in chapters:
            document.add_heading(f"Chapter {ch['number']}: {ch['title']}", level=1)
            for block in self._iter_render_blocks(ch["content"]):
                block_type = str(block.get("type", "paragraph"))
                block_text = str(block.get("text", "")).strip()
                if not block_text and block_type not in {"visual_placeholder"}:
                    continue
                if block_type == "h1":
                    document.add_heading(block_text, level=1)
                elif block_type == "h2":
                    document.add_heading(block_text, level=2)
                elif block_type == "code":
                    p = document.add_paragraph()
                    run = p.add_run(block_text)
                    run.font.name = "Courier New"
                    run.font.size = Pt(9)
                elif block_type == "quote":
                    p = document.add_paragraph()
                    run = p.add_run(block_text)
                    run.italic = True
                elif block_type == "callout":
                    p = document.add_paragraph()
                    label = str(block.get("label", "Note")).strip() or "Note"
                    run = p.add_run(f"{label}: {block_text}")
                    run.italic = True
                elif block_type == "visual_placeholder":
                    kind = str(block.get("kind", "visual")).upper()
                    label = str(block.get("label", "")).strip()
                    p = document.add_paragraph()
                    run = p.add_run(f"[{kind} PLACEHOLDER] {label}")
                    run.italic = True
                else:
                    document.add_paragraph(block_text)
            document.add_page_break()

        out = io.BytesIO()
        document.save(out)
        out.seek(0)
        return out.read()

    def _split_blocks(self, text: str) -> List[str]:
        if not text:
            return []
        norm = text.replace("\r\n", "\n").replace("\r", "\n")
        blocks = re.split(r"\n\s*\n", norm)
        return [b.strip() for b in blocks if b.strip()]

    def _has_visual_placeholders(self, text: str) -> bool:
        return bool(re.search(r"(?im)^\[(FIGURE|FLOWCHART)\s*:", str(text or "")))

    def _iter_render_blocks(self, text: str) -> List[Dict[str, str]]:
        norm = str(text or "").replace("\r\n", "\n").replace("\r", "\n")
        if not norm.strip():
            return []

        lines = norm.split("\n")
        blocks: List[Dict[str, str]] = []
        i = 0
        while i < len(lines):
            line = lines[i]
            stripped = line.strip()
            if not stripped:
                i += 1
                continue

            if stripped.startswith("```"):
                code_lines: List[str] = []
                i += 1
                while i < len(lines) and not lines[i].strip().startswith("```"):
                    code_lines.append(lines[i])
                    i += 1
                if i < len(lines) and lines[i].strip().startswith("```"):
                    i += 1
                blocks.append({"type": "code", "text": "\n".join(code_lines).rstrip()})
                continue

            visual_match = re.match(r"^\[(FIGURE|FLOWCHART)\s*:\s*(.+?)\]\s*$", stripped, flags=re.IGNORECASE)
            if visual_match:
                blocks.append(
                    {
                        "type": "visual_placeholder",
                        "kind": visual_match.group(1).strip().lower(),
                        "label": visual_match.group(2).strip(),
                        "text": "",
                    }
                )
                i += 1
                continue

            if stripped.startswith(">"):
                quote_lines: List[str] = []
                callout_label = ""
                while i < len(lines) and lines[i].strip().startswith(">"):
                    raw_line = lines[i].strip()[1:].lstrip()
                    callout_match = re.match(r"^\[\!([A-Z]+)\]\s*(.*)$", raw_line)
                    if callout_match:
                        callout_label = callout_match.group(1).title()
                        raw_line = callout_match.group(2).strip()
                    quote_lines.append(raw_line)
                    i += 1
                blocks.append(
                    {
                        "type": "callout" if callout_label else "quote",
                        "label": callout_label,
                        "text": "\n".join([line for line in quote_lines if line]).strip(),
                    }
                )
                continue

            if stripped.startswith("# "):
                blocks.append({"type": "h1", "text": stripped[2:].strip()})
                i += 1
                continue
            if stripped.startswith("## "):
                blocks.append({"type": "h2", "text": stripped[3:].strip()})
                i += 1
                continue

            if "|" in stripped:
                table_lines: List[str] = []
                start = i
                while i < len(lines):
                    probe = lines[i].strip()
                    if not probe or "|" not in probe:
                        break
                    table_lines.append(probe)
                    i += 1
                table_text = "\n".join(table_lines).strip()
                if len(table_lines) >= 2 and re.search(r"(?m)^\|?\s*:?-{2,}", table_text):
                    blocks.append({"type": "table", "text": table_text})
                    continue
                i = start

            paragraph_lines: List[str] = []
            while i < len(lines):
                probe = lines[i]
                probe_stripped = probe.strip()
                if not probe_stripped:
                    break
                if probe_stripped.startswith(("```", "# ", "## ", ">", "[")):
                    break
                paragraph_lines.append(probe_stripped)
                i += 1
            if paragraph_lines:
                blocks.append({"type": "paragraph", "text": " ".join(paragraph_lines).strip()})
            else:
                i += 1

        return blocks

    def _escape(self, value: str) -> str:
        return value.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")

    def _safe_file(self, title: str) -> str:
        sanitized = re.sub(r"[<>:\"/\\|?*]", "_", title).strip().strip(".")
        sanitized = sanitized.replace(" ", "_")
        return sanitized[:100] or "book"

    def _merge_project_metadata(self, project: BookProject, llm_metadata: Dict[str, Any] | Any) -> Dict[str, Any]:
        existing = project.metadata_json if isinstance(project.metadata_json, dict) else {}
        user_concept = self._build_user_concept_snapshot(project, existing)
        llm_runtime = llm_metadata if isinstance(llm_metadata, dict) else {}
        merged = dict(existing)
        merged["user_concept"] = user_concept
        merged["llm_runtime"] = llm_runtime
        if isinstance(user_concept.get("profile"), dict):
            merged["profile"] = user_concept["profile"]
        if "subtitle" in user_concept:
            merged["subtitle"] = user_concept["subtitle"]
        if "instruction_brief" in user_concept:
            merged["instruction_brief"] = user_concept["instruction_brief"]
        return merged

    def _outline_profile_compliance(self, project: BookProject, outline: Dict[str, Any]) -> Dict[str, Any]:
        profile = self._project_profile(project)
        checks: Dict[str, Any] = {}
        issues: List[str] = []

        chapters = outline.get("chapters", []) if isinstance(outline, dict) else []
        if not isinstance(chapters, list):
            chapters = []
        chapter_count = len(chapters)
        checks["chapter_count"] = chapter_count

        empty_bullet_chapters: List[int] = []
        for chapter in chapters:
            if not isinstance(chapter, dict):
                continue
            bullet_points = chapter.get("bullet_points", [])
            if not isinstance(bullet_points, list):
                empty_bullet_chapters.append(int(chapter.get("number", 0)))
                continue
            non_empty_points = [bp for bp in bullet_points if str(bp).strip()]
            if not non_empty_points:
                empty_bullet_chapters.append(int(chapter.get("number", 0)))
        if empty_bullet_chapters:
            issues.append(
                f"Outline compliance: chapters without bullet points: {', '.join(str(n) for n in empty_bullet_chapters)}."
            )
            checks["empty_bullet_chapters"] = empty_bullet_chapters

        chapter_length = str(profile.get("chapterLength", "")).strip().lower()
        words_per_chapter = 0
        if "short" in chapter_length:
            words_per_chapter = 1500
        elif "medium" in chapter_length:
            words_per_chapter = 3000
        elif "long" in chapter_length:
            words_per_chapter = 5000

        target_word_count = 0
        try:
            target_word_count = max(0, int(float(str(profile.get("length", project.target_word_count) or 0))))
        except Exception:
            target_word_count = max(0, int(project.target_word_count or 0))

        if words_per_chapter > 0 and target_word_count > 0 and chapter_count > 0:
            expected_chapters = max(1, round(target_word_count / words_per_chapter))
            deviation = abs(chapter_count - expected_chapters)
            deviation_ratio = deviation / max(1, expected_chapters)
            checks["chapter_count_vs_length"] = {
                "target_word_count": target_word_count,
                "chapter_length": str(profile.get("chapterLength", "")).strip(),
                "expected_chapters": expected_chapters,
                "actual_chapters": chapter_count,
                "deviation": deviation,
                "deviation_ratio": round(deviation_ratio, 2),
            }
            if deviation_ratio > 0.6 and deviation >= 2:
                issues.append(
                    "Outline compliance: chapter count may not match the selected chapter length and total word count."
                )

        return {
            "fail": bool(issues),
            "issues": issues,
            "checks": checks,
        }

    def _analyze_refine_feedback_conflicts(self, project: BookProject, feedback: str) -> Dict[str, Any]:
        profile = self._project_profile(project)
        text = str(feedback or "").strip()
        normalized = text.lower()
        issues: List[str] = []
        checks: Dict[str, Any] = {}

        point_of_view = str(profile.get("pointOfView", "")).strip()
        pov_lower = point_of_view.lower()
        if point_of_view:
            requested_pov: str | None = None
            if re.search(r"\b(first[\s-]?person|memoir|i perspective)\b", normalized):
                requested_pov = "First Person"
            elif re.search(r"\b(second[\s-]?person|you perspective)\b", normalized):
                requested_pov = "Second Person"
            elif re.search(r"\b(third[\s-]?person)\b", normalized):
                requested_pov = "Third Person"
            if requested_pov:
                checks["pointOfView"] = {"profile": point_of_view, "feedback_request": requested_pov}
                if (
                    ("first" in pov_lower and requested_pov != "First Person")
                    or ("second" in pov_lower and requested_pov != "Second Person")
                    or ("third" in pov_lower and requested_pov != "Third Person")
                ):
                    issues.append(
                        f"Refine feedback may conflict with saved point of view ({point_of_view}) by requesting {requested_pov.lower()}."
                    )

        audience_level = str(profile.get("audienceKnowledgeLevel", "")).strip()
        vocabulary_level = str(profile.get("vocabularyLevel", "")).strip()
        beginner_target = "beginner" in audience_level.lower() or vocabulary_level.lower() == "simple"
        if beginner_target:
            if re.search(r"\b(expert|advanced|highly technical|assume prior knowledge|graduate level)\b", normalized):
                checks["readability"] = {
                    "profile_audience_level": audience_level,
                    "profile_vocabulary_level": vocabulary_level,
                    "feedback": "more advanced/technical",
                }
                issues.append(
                    "Refine feedback may conflict with the beginner/simple readability target by asking for advanced technical depth."
                )

        tone = str(profile.get("tone", "")).strip()
        tone_lower = tone.lower()
        if tone:
            requested_tone: str | None = None
            if re.search(r"\b(formal|academic)\b", normalized):
                requested_tone = "Formal/Academic"
            elif re.search(r"\b(conversational|casual|friendly)\b", normalized):
                requested_tone = "Conversational"
            elif re.search(r"\b(humorous|funny|playful)\b", normalized):
                requested_tone = "Humorous/Playful"
            if requested_tone:
                checks["tone"] = {"profile": tone, "feedback_signal": requested_tone}
                if (
                    ("formal" in tone_lower or "academic" in tone_lower) and requested_tone in {"Conversational", "Humorous/Playful"}
                ) or (
                    ("conversational" in tone_lower or "informative" in tone_lower) and requested_tone == "Formal/Academic"
                ):
                    issues.append(f"Refine feedback may conflict with saved tone ({tone}).")

        content_boundaries = str(profile.get("contentBoundaries", "")).strip()
        if content_boundaries:
            if re.search(r"\b(ignore|remove|drop|relax|bypass|override)\b.{0,40}\b(boundar|restriction|safety|limit)\b", normalized):
                checks["contentBoundaries"] = {"profile": content_boundaries[:240], "feedback_signal": "weaken/remove boundaries"}
                issues.append("Refine feedback appears to weaken saved content boundaries. Review before applying.")

        book_purpose = str(profile.get("bookPurpose", "")).strip()
        if book_purpose and re.search(r"\b(turn this into|change purpose to|make it a story instead|make it a research report)\b", normalized):
            checks["bookPurpose"] = {"profile": book_purpose}
            issues.append(f"Refine feedback may change the saved book purpose ({book_purpose}).")

        return {
            "warn": bool(issues),
            "issues": issues,
            "checks": checks,
        }

    def _project_profile(self, project: BookProject) -> Dict[str, Any]:
        metadata = project.metadata_json if isinstance(project.metadata_json, dict) else {}
        user_concept = metadata.get("user_concept", {})
        if isinstance(user_concept, dict):
            profile = user_concept.get("profile", {})
            if isinstance(profile, dict):
                return profile
        legacy_profile = metadata.get("profile", {})
        return legacy_profile if isinstance(legacy_profile, dict) else {}

    def _build_user_concept_snapshot(self, project: BookProject, existing_meta: Dict[str, Any]) -> Dict[str, Any]:
        existing_user = existing_meta.get("user_concept", {})
        user_concept = dict(existing_user) if isinstance(existing_user, dict) else {}
        user_concept.setdefault("title", project.title)
        user_concept.setdefault("genre", project.genre)
        user_concept.setdefault("target_audience", project.target_audience)
        user_concept.setdefault("language", project.language)
        user_concept.setdefault("tone", project.tone)
        user_concept.setdefault("target_word_count", project.target_word_count)

        if not isinstance(user_concept.get("profile"), dict):
            legacy_profile = existing_meta.get("profile", {})
            user_concept["profile"] = legacy_profile if isinstance(legacy_profile, dict) else {}
        if "subtitle" not in user_concept and isinstance(existing_meta.get("subtitle"), str):
            user_concept["subtitle"] = existing_meta.get("subtitle", "")
        if "instruction_brief" not in user_concept and isinstance(existing_meta.get("instruction_brief"), str):
            user_concept["instruction_brief"] = existing_meta.get("instruction_brief", "")
        return user_concept
