"""
Book Writing Agent  Fully Deployable, UI-Friendly, AgentGrid-Compliant (Stateless)
File: backend/agentgrid-backend/app/agents/book_writing_agent.py

What this agent supports (industry workflow):
1) mode="toc"          -> Generate Table of Contents (outline) first (human-in-loop)
2) mode="refine_toc"   -> Refine an existing outline based on user feedback/edits
3) mode="chapter"      -> Generate a single chapter (chapter_number) using the approved outline
4) mode="export"       -> Export a real book layout to PDF and/or DOCX using provided outline + chapters

IMPORTANT (Stateless design):
- The agent does not store anything.
- Your UI/backend should store the outline and generated chapters.
- For export, your UI sends back the approved outline + chapters content.

Exports:
- PDF: cover page, TOC, chapter page breaks, headings, page numbers
- DOCX: cover, headings, chapters; TOC is inserted as a placeholder note (Word can auto-generate TOC from headings)

Env Config:
- OPENAI_API_KEY (required)
- BOOK_AGENT_MODEL (default: gpt-4o)
- BOOK_AGENT_TEMPERATURE (default: 0.7)
- BOOK_AGENT_MAX_TOKENS (default: 4096)
- BOOK_AGENT_TIMEOUT_S (default: 90)
- BOOK_AGENT_JSON_RETRIES (default: 2)
- BOOK_AGENT_MAX_BOOK_LENGTH (default: 200000)
- BOOK_AGENT_MIN_BOOK_LENGTH (default: 300)
- BOOK_AGENT_MAX_PDF_BYTES (default: 8MB)
- BOOK_AGENT_MAX_DOCX_BYTES (default: 8MB)
"""

from __future__ import annotations

from typing import Dict, Any, List, Optional
import os
import json
import base64
import io
import logging
import time
import uuid
import re
from dotenv import load_dotenv

from app.agents.base import BaseAgent, AgentInput, AgentOutput
from app.agents.registry import register_agent

from langchain_openai import ChatOpenAI
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.messages import SystemMessage, HumanMessage

# PDF (ReportLab)
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch, mm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak

# DOCX (python-docx)
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)

BOOK_AGENT_ID = "eef314c9-183b-4d87-9d6c-88815a72be15"

DOTENV_LOADED = False

# -----------------------------
# Configuration (env override)
# -----------------------------
DEFAULT_MODEL = os.getenv("BOOK_AGENT_MODEL", "gpt-4o")
DEFAULT_TEMPERATURE = float(os.getenv("BOOK_AGENT_TEMPERATURE", "0.7"))
DEFAULT_MAX_TOKENS = int(os.getenv("BOOK_AGENT_MAX_TOKENS", "4096"))
DEFAULT_TIMEOUT_S = int(os.getenv("BOOK_AGENT_TIMEOUT_S", "90"))

JSON_PARSE_RETRIES = int(os.getenv("BOOK_AGENT_JSON_RETRIES", "2"))

MAX_BOOK_LENGTH_WORDS = int(os.getenv("BOOK_AGENT_MAX_BOOK_LENGTH", "200000"))
MIN_BOOK_LENGTH_WORDS = int(os.getenv("BOOK_AGENT_MIN_BOOK_LENGTH", "300"))

MAX_PDF_BYTES = int(os.getenv("BOOK_AGENT_MAX_PDF_BYTES", str(8 * 1024 * 1024)))
MAX_DOCX_BYTES = int(os.getenv("BOOK_AGENT_MAX_DOCX_BYTES", str(8 * 1024 * 1024)))

MAX_TITLE_LEN = 160
MAX_GENRE_LEN = 80
MAX_TONE_LEN = 80
MAX_LANGUAGE_LEN = 40
MAX_AUDIENCE_LEN = 80
MAX_FEEDBACK_LEN = 2000


@register_agent(BOOK_AGENT_ID)
class BookWritingAgent(BaseAgent):
    # -----------------
    # AgentGrid metadata
    # -----------------
    @property
    def name(self) -> str:
        return "Book Writing Agent"

    @property
    def description(self) -> str:
        return (
            "Human-in-the-loop book generator: TOC -> refine -> chapter -> export (PDF/DOCX). "
            "Stateless, UI-friendly, production safe."
        )

    @property
    def inputs(self) -> List[AgentInput]:
        return [
            AgentInput("mode", "string", "toc | refine_toc | chapter | export"),
            AgentInput("book_title", "string", "Title of the book"),
            AgentInput("genre", "string", "Genre (e.g., Fiction, Non-fiction, Educational)"),
            AgentInput("target_audience", "string", "Target audience (optional)", required=False),
            AgentInput("language", "string", "Language (e.g., English)"),
            AgentInput("tone", "string", "Tone (e.g., friendly, academic)"),
            AgentInput("book_length", "number", "Target word count"),
            AgentInput("model", "string", "Override model (optional)", required=False),
            AgentInput("outline", "string", "Approved outline JSON (required for refine_toc/chapter/export)", required=False),
            AgentInput("feedback", "string", "User feedback for refining TOC (refine_toc mode)", required=False),
            AgentInput("chapter_number", "number", "Chapter number to generate (chapter mode)", required=False),
            AgentInput("chapters", "string", "Chapters JSON for export: [{number,title,content}, ...]", required=False),
            AgentInput("export_format", "string", "pdf | docx | both (export mode)", required=False, options=["pdf", "docx", "both"]),
        ]

    @property
    def outputs(self) -> List[AgentOutput]:
        # Keep outputs stable for UI
        return [
            AgentOutput("status", "string", "success|error"),
            AgentOutput("trace_id", "string", "Trace id for logs/support"),
            AgentOutput("errors", "array", "User-safe errors"),
            AgentOutput("warnings", "array", "Non-fatal warnings"),
            AgentOutput("timings_ms", "object", "Timing metrics"),

            AgentOutput("outline", "object", "Outline/TOC object"),
            AgentOutput("chapter", "object", "Single chapter object (chapter mode)"),
            AgentOutput("metadata", "object", "Metadata (themes, keywords, etc.)"),
            AgentOutput("next_steps", "array", "Suggested next actions"),

            AgentOutput("pdf_base64", "string", "PDF bytes base64 (export mode)"),
            AgentOutput("pdf_filename", "string", "PDF filename"),
            AgentOutput("docx_base64", "string", "DOCX bytes base64 (export mode)"),
            AgentOutput("docx_filename", "string", "DOCX filename"),
        ]

    # -------------
    # Main execution
    # -------------
    def run(self, inputs: Dict[str, Any]) -> Dict[str, Any]:
        trace_id = str(uuid.uuid4())
        t0 = time.perf_counter()

        result: Dict[str, Any] = {
            "status": "error",
            "trace_id": trace_id,
            "errors": [],
            "warnings": [],
            "timings_ms": {},

            "outline": {"synopsis": "", "chapters": []},
            "chapter": None,
            "metadata": {},
            "next_steps": [],

            "pdf_base64": None,
            "pdf_filename": None,
            "docx_base64": None,
            "docx_filename": None,
        }

        try:
            mode = self._clean_str(inputs.get("mode"), 32) or ""
            if mode not in {"toc", "refine_toc", "chapter", "export"}:
                raise ValueError("mode must be one of: toc | refine_toc | chapter | export")

            # Normalize common requirements (toc/refine/chapter/export)
            common = self._normalize_common_inputs(inputs, result["warnings"], mode)
            result["timings_ms"]["validate"] = self._ms_since(t0)

            llm = None
            if mode in {"toc", "refine_toc", "chapter"}:
                llm = self._initialize_llm(inputs)

            if mode == "toc":
                t_llm = time.perf_counter()
                payload = self._generate_toc(llm, common)
                result["timings_ms"]["llm"] = self._ms_since(t_llm)
                result["outline"] = payload["outline"]
                result["metadata"] = payload.get("metadata", {})
                result["next_steps"] = payload.get("next_steps", [
                    "Review the Table of Contents (TOC).",
                    "Edit chapter titles/order if needed.",
                    "Run refine_toc if you want the agent to improve the TOC based on your feedback.",
                    "Generate chapters one by one using chapter mode."
                ])
                result["status"] = "success"

            elif mode == "refine_toc":
                outline = self._require_outline(inputs)
                feedback = self._clean_str(inputs.get("feedback"), MAX_FEEDBACK_LEN)
                if not feedback:
                    raise ValueError("feedback is required for refine_toc mode")

                t_llm = time.perf_counter()
                payload = self._refine_toc(llm, common, outline, feedback)
                result["timings_ms"]["llm"] = self._ms_since(t_llm)

                result["outline"] = payload["outline"]
                result["metadata"] = payload.get("metadata", {})
                result["next_steps"] = payload.get("next_steps", [
                    "Review the refined TOC.",
                    "If satisfied, start generating chapters using chapter mode.",
                    "If not satisfied, refine again with more specific feedback."
                ])
                result["status"] = "success"

            elif mode == "chapter":
                outline = self._require_outline(inputs)
                chapter_number = self._to_int_required(inputs.get("chapter_number"), "chapter_number")
                self._validate_chapter_number_against_outline(chapter_number, outline)

                t_llm = time.perf_counter()
                payload = self._generate_chapter(llm, common, outline, chapter_number)
                result["timings_ms"]["llm"] = self._ms_since(t_llm)

                result["outline"] = outline  # echo for UI convenience
                result["chapter"] = payload["chapter"]
                result["metadata"] = payload.get("metadata", {})
                result["next_steps"] = payload.get("next_steps", [
                    "Review the chapter content.",
                    "If you want changes, regenerate this chapter with refined feedback in your UI (e.g., add a feedback field).",
                    "Generate the next chapter when ready.",
                    "When all chapters are finalized, use export mode to create PDF/DOCX."
                ])
                result["status"] = "success"

            else:  # export
                outline = self._require_outline(inputs)
                chapters = self._require_chapters(inputs)
                export_format = self._clean_str(inputs.get("export_format"), 16).lower() or "pdf"
                if export_format not in {"pdf", "docx", "both"}:
                    raise ValueError("export_format must be one of: pdf | docx | both")

                # Validate that chapters include required fields and are ordered
                chapters_norm = self._normalize_export_chapters(chapters)
                result["outline"] = outline

                t_export = time.perf_counter()
                if export_format in {"pdf", "both"}:
                    pdf_bytes = self._generate_book_pdf(
                        title=common["book_title"],
                        outline=outline,
                        chapters=chapters_norm,
                        metadata=common,
                    )
                    if len(pdf_bytes) > MAX_PDF_BYTES:
                        result["warnings"].append(
                            f"PDF exceeds size limit ({len(pdf_bytes)} bytes). Try exporting fewer chapters or shorten content."
                        )
                    else:
                        result["pdf_base64"] = base64.b64encode(pdf_bytes).decode("utf-8")
                        result["pdf_filename"] = f"{self._sanitize_filename(common['book_title'])}.pdf"

                if export_format in {"docx", "both"}:
                    docx_bytes = self._generate_book_docx(
                        title=common["book_title"],
                        outline=outline,
                        chapters=chapters_norm,
                        metadata=common,
                    )
                    if len(docx_bytes) > MAX_DOCX_BYTES:
                        result["warnings"].append(
                            f"DOCX exceeds size limit ({len(docx_bytes)} bytes). Try exporting fewer chapters or shorten content."
                        )
                    else:
                        result["docx_base64"] = base64.b64encode(docx_bytes).decode("utf-8")
                        result["docx_filename"] = f"{self._sanitize_filename(common['book_title'])}.docx"

                result["timings_ms"]["export"] = self._ms_since(t_export)
                result["next_steps"] = [
                    "Download the exported file(s).",
                    "For DOCX: open in Word/Google Docs and (optionally) insert/update a Table of Contents from headings.",
                    "Publish or share your final book."
                ]
                result["status"] = "success"

            result["timings_ms"]["total"] = self._ms_since(t0)
            logger.info("BookWritingAgent done mode=%s trace_id=%s status=%s", mode, trace_id, result["status"])
            return result

        except ValueError as e:
            msg = str(e) if str(e) else "Invalid request."
            result["errors"].append(msg)
            result["timings_ms"]["total"] = self._ms_since(t0)
            logger.warning("BookWritingAgent value_error trace_id=%s err=%s", trace_id, msg)
            return result

        except Exception:
            # Keep user-safe; log full stack trace
            result["errors"].append("Agent execution failed. Please try again.")
            result["timings_ms"]["total"] = self._ms_since(t0)
            logger.error("BookWritingAgent unexpected trace_id=%s", trace_id, exc_info=True)
            return result

    # --------------------
    # Common input handling
    # --------------------
    def _normalize_common_inputs(self, inputs: Dict[str, Any], warnings: List[str], mode: str) -> Dict[str, Any]:
        # For export, we still require common fields so filenames and cover look correct.
        book_title = self._clean_required_str(inputs.get("book_title"), "book_title", MAX_TITLE_LEN)
        genre = self._clean_required_str(inputs.get("genre"), "genre", MAX_GENRE_LEN)
        language = self._clean_required_str(inputs.get("language"), "language", MAX_LANGUAGE_LEN)
        tone = self._clean_required_str(inputs.get("tone"), "tone", MAX_TONE_LEN)

        target_audience = self._clean_str(inputs.get("target_audience"), MAX_AUDIENCE_LEN) or "General readers"

        # book_length required for toc/refine/chapter; for export allow missing if you already have chapters
        book_length_raw = inputs.get("book_length")
        if mode in {"toc", "refine_toc", "chapter"}:
            book_length = self._to_int_required(book_length_raw, "book_length")
        else:
            book_length = self._to_int_optional(book_length_raw) or 0

        if book_length:
            if book_length < MIN_BOOK_LENGTH_WORDS:
                warnings.append(f"book_length increased to minimum {MIN_BOOK_LENGTH_WORDS}.")
                book_length = MIN_BOOK_LENGTH_WORDS
            if book_length > MAX_BOOK_LENGTH_WORDS:
                raise ValueError(f"book_length cannot exceed {MAX_BOOK_LENGTH_WORDS} words")

        return {
            "book_title": book_title,
            "genre": genre,
            "target_audience": target_audience,
            "language": language,
            "tone": tone,
            "book_length": book_length,
        }

    # ----------------
    # LLM initialization
    # ----------------
    def _initialize_llm(self, inputs: Dict[str, Any]) -> ChatOpenAI:
        global DOTENV_LOADED
        if not DOTENV_LOADED:
            load_dotenv(override=False)
            DOTENV_LOADED = True

        api_key = os.getenv("OPENAI_API_KEY")
        if not api_key:
            raise ValueError("OPENAI_API_KEY environment variable not set")

        model = self._clean_str(inputs.get("model"), 64) or DEFAULT_MODEL
        return ChatOpenAI(
            api_key=api_key,
            model=model,
            temperature=DEFAULT_TEMPERATURE,
            max_tokens=DEFAULT_MAX_TOKENS,
            request_timeout=DEFAULT_TIMEOUT_S,
        )

    # ----------------
    # TOC generation
    # ----------------
    def _generate_toc(self, llm: ChatOpenAI, common: Dict[str, Any]) -> Dict[str, Any]:
        system_msg = (
            "You are a professional book architect.\n"
            "Return ONLY valid JSON. No markdown. No code fences.\n\n"
            "Schema:\n"
            "{\n"
            '  "outline": {\n'
            '    "synopsis": "string",\n'
            '    "chapters": [\n'
            '      {"number": 1, "title": "string", "bullet_points": ["string", "..."]}\n'
            '    ]\n'
            '  },\n'
            '  "metadata": {"key_themes": ["..."], "seo_keywords": ["..."]},\n'
            '  "next_steps": ["..."]\n'
            "}\n"
        )

        human_msg = (
            "Create a high-quality Table of Contents (TOC) and synopsis.\n"
            f"Title: {common['book_title']}\n"
            f"Genre: {common['genre']}\n"
            f"Audience: {common['target_audience']}\n"
            f"Language: {common['language']}\n"
            f"Tone: {common['tone']}\n"
            f"Target Word Count: {common['book_length']}\n\n"
            "Guidelines:\n"
            "- Use a logical chapter progression.\n"
            "- Provide 36 bullet points per chapter.\n"
            "- Keep it realistic for the target word count.\n"
            "- Output valid JSON only.\n"
        )

        payload = self._llm_json_with_retries(llm, system_msg, human_msg)
        self._validate_outline(payload.get("outline"))
        return payload

    # ----------------
    # TOC refinement
    # ----------------
    def _refine_toc(self, llm: ChatOpenAI, common: Dict[str, Any], outline: Dict[str, Any], feedback: str) -> Dict[str, Any]:
        system_msg = (
            "You are a professional editor improving a book Table of Contents.\n"
            "Return ONLY valid JSON. No markdown. No code fences.\n\n"
            "Schema:\n"
            "{\n"
            '  "outline": {\n'
            '    "synopsis": "string",\n'
            '    "chapters": [{"number": 1, "title": "string", "bullet_points": ["..."]}]\n'
            '  },\n'
            '  "metadata": {"key_themes": ["..."], "seo_keywords": ["..."]},\n'
            '  "next_steps": ["..."]\n'
            "}\n"
        )

        human_msg = (
            "Refine the existing TOC based on the user's feedback.\n\n"
            f"Book Title: {common['book_title']}\n"
            f"Genre: {common['genre']}\n"
            f"Audience: {common['target_audience']}\n"
            f"Language: {common['language']}\n"
            f"Tone: {common['tone']}\n"
            f"Target Word Count: {common['book_length']}\n\n"
            f"User Feedback:\n{feedback}\n\n"
            f"Existing Outline (JSON):\n{json.dumps(outline, ensure_ascii=False)}\n\n"
            "Rules:\n"
            "- Keep chapter numbering sequential starting at 1.\n"
            "- Improve structure and clarity.\n"
            "- Keep TOC coherent with target length.\n"
            "- Output valid JSON only.\n"
        )

        payload = self._llm_json_with_retries(llm, system_msg, human_msg)
        self._validate_outline(payload.get("outline"))
        return payload

        payload = self._llm_json_with_retries(llm, system_msg, human_msg)
        self._validate_outline(payload.get("outline"))
        return payload

    # ----------------
    # Chapter generation
    # ----------------
    def _generate_chapter(self, llm: ChatOpenAI, common: Dict[str, Any], outline: Dict[str, Any], chapter_number: int) -> Dict[str, Any]:
        chapter_title = self._get_chapter_title_from_outline(outline, chapter_number)

        system_msg = (
            "You are a professional author.\n"
            "Return ONLY valid JSON. No markdown. No code fences.\n\n"
            "Schema:\n"
            "{\n"
            '  "chapter": {"number": 1, "title": "string", "content": "string", "summary": "string"},\n'
            '  "metadata": {"key_themes": ["..."], "seo_keywords": ["..."]},\n'
            '  "next_steps": ["..."]\n'
            "}\n"
            "Content formatting rules:\n"
            "- Use plain text with simple headings inside content like:\n"
            "  # Main Heading\n"
            "  ## Subheading\n"
            "- Keep paragraphs separated by blank lines.\n"
        )

        human_msg = (
            "Write ONE chapter for the approved outline.\n\n"
            f"Book Title: {common['book_title']}\n"
            f"Genre: {common['genre']}\n"
            f"Audience: {common['target_audience']}\n"
            f"Language: {common['language']}\n"
            f"Tone: {common['tone']}\n"
            f"Target Word Count (whole book): {common['book_length']}\n\n"
            f"Approved Outline (JSON):\n{json.dumps(outline, ensure_ascii=False)}\n\n"
            f"Chapter to write:\n- Number: {chapter_number}\n- Title: {chapter_title}\n\n"
            "Rules:\n"
            "- Return valid JSON only.\n"
            "- Keep the chapter coherent with the outline.\n"
            "- Make headings clear using '# ' and '## ' inside content.\n"
        )

        payload = self._llm_json_with_retries(llm, system_msg, human_msg)
        chap = payload.get("chapter")
        if not isinstance(chap, dict):
            raise ValueError("LLM response missing chapter object")
        if chap.get("number") != chapter_number:
            # normalize number if model changes it
            chap["number"] = chapter_number
        if not chap.get("title"):
            chap["title"] = chapter_title
        if not isinstance(chap.get("content"), str) or not chap["content"].strip():
            raise ValueError("Generated chapter content is empty")
        return payload

    # -----------------------------
    # LLM JSON helper with retries
    # -----------------------------
    def _llm_json_with_retries(self, llm: ChatOpenAI, system_msg: str, human_msg: str) -> Dict[str, Any]:
        messages = [
            SystemMessage(content=system_msg),
            HumanMessage(content=human_msg)
        ]

        last_err: Optional[str] = None
        for attempt in range(JSON_PARSE_RETRIES + 1):
            resp = llm.invoke(messages)
            content = (resp.content or "").strip()
            try:
                data = self._extract_json_robust(content)
                if not isinstance(data, dict):
                    raise ValueError("Top-level JSON must be an object")
                return data
            except Exception as e:
                last_err = str(e) or "invalid_json"
                logger.warning("JSON parse failed attempt=%s err=%s", attempt + 1, last_err)

                # Ask model to fix JSON only (no new content)
                fix_msg = (
                    "Your previous response was not valid JSON or did not match the schema.\n"
                    "Return ONLY corrected valid JSON that matches the schema exactly.\n"
                    "No markdown, no code fences, no explanations."
                )
                messages.append(HumanMessage(content=fix_msg))

        raise ValueError("LLM returned invalid JSON repeatedly. Please try again.")

    def _extract_json_robust(self, content: str) -> Dict[str, Any]:
        c = content.strip()

        # Strip fenced blocks if present
        if "```" in c:
            parts = c.split("```")
            for mid in parts[1:-1]:
                candidate = mid.strip()
                # remove optional language tag
                lines = candidate.splitlines()
                if lines and lines[0].strip().lower() in ("json", "javascript", "js"):
                    candidate = "\n".join(lines[1:]).strip()
                try:
                    return json.loads(candidate)
                except Exception:
                    pass

        # Fallback: outermost JSON object
        start = c.find("{")
        end = c.rfind("}") + 1
        if start == -1 or end <= start:
            raise ValueError("No JSON object found")
        return json.loads(c[start:end])

    # -----------------------
    # Outline/chapter validation
    # -----------------------
    def _require_outline(self, inputs: Dict[str, Any]) -> Dict[str, Any]:
        outline_raw = inputs.get("outline")
        outline = self._parse_json_input(outline_raw, "outline")
        if not isinstance(outline, dict):
            raise ValueError("outline is required and must be an object for this mode")
        self._validate_outline(outline)
        return outline

    def _validate_outline(self, outline: Any) -> None:
        if not isinstance(outline, dict):
            raise ValueError("outline must be an object")
        chapters = outline.get("chapters")
        if not isinstance(chapters, list) or not chapters:
            raise ValueError("outline.chapters must be a non-empty array")
        # ensure sequential numbering
        expected = 1
        for ch in chapters:
            if not isinstance(ch, dict):
                raise ValueError("Each outline chapter must be an object")
            if ch.get("number") != expected:
                raise ValueError("Outline chapter numbers must be sequential starting at 1")
            if not str(ch.get("title", "")).strip():
                raise ValueError(f"Outline chapter {expected} missing title")
            bps = ch.get("bullet_points", [])
            if bps is not None and not isinstance(bps, list):
                raise ValueError(f"Outline chapter {expected} bullet_points must be an array")
            expected += 1

    def _validate_chapter_number_against_outline(self, chapter_number: int, outline: Dict[str, Any]) -> None:
        total = len(outline.get("chapters", []))
        if chapter_number < 1 or chapter_number > total:
            raise ValueError(f"chapter_number must be between 1 and {total} (based on outline)")

    def _get_chapter_title_from_outline(self, outline: Dict[str, Any], chapter_number: int) -> str:
        for ch in outline.get("chapters", []):
            if ch.get("number") == chapter_number:
                return str(ch.get("title")).strip()
        return f"Chapter {chapter_number}"

    # -----------------------
    # Export input normalization
    # -----------------------
    def _require_chapters(self, inputs: Dict[str, Any]) -> List[Dict[str, Any]]:
        chapters_raw = inputs.get("chapters")
        chapters = self._parse_json_input(chapters_raw, "chapters")
        if not isinstance(chapters, list) or not chapters:
            raise ValueError("chapters is required and must be a non-empty array for export mode")
        return chapters

    def _normalize_export_chapters(self, chapters: List[Any]) -> List[Dict[str, Any]]:
        out: List[Dict[str, Any]] = []
        for item in chapters:
            if not isinstance(item, dict):
                raise ValueError("Each chapters[] item must be an object {number,title,content}")
            num = item.get("number")
            title = str(item.get("title", "")).strip()
            content = item.get("content", "")
            if num is None:
                raise ValueError("Each chapter must have a number")
            try:
                num_int = int(num)
            except Exception:
                raise ValueError("Chapter number must be an integer")
            if not title:
                title = f"Chapter {num_int}"
            if not isinstance(content, str) or not content.strip():
                raise ValueError(f"Chapter {num_int} content is empty")
            out.append({"number": num_int, "title": title, "content": content})
        out.sort(key=lambda x: x["number"])
        # ensure sequential
        expected = 1
        for ch in out:
            if ch["number"] != expected:
                raise ValueError("Export chapters must be sequential starting at 1 (no gaps)")
            expected += 1
        return out

    # -----------------------
    # PDF export: real book layout
    # -----------------------
    def _generate_book_pdf(
        self,
        title: str,
        outline: Dict[str, Any],
        chapters: List[Dict[str, Any]],
        metadata: Dict[str, Any],
    ) -> bytes:
        buf = io.BytesIO()
        doc = SimpleDocTemplate(
            buf,
            pagesize=A4,
            topMargin=1 * inch,
            bottomMargin=1 * inch,
            leftMargin=1 * inch,
            rightMargin=1 * inch,
        )

        styles = getSampleStyleSheet()
        # Custom styles (safe, readable)
        h1 = ParagraphStyle("BookH1", parent=styles["Heading1"], spaceAfter=12)
        h2 = ParagraphStyle("BookH2", parent=styles["Heading2"], spaceAfter=10)
        body = ParagraphStyle("BookBody", parent=styles["BodyText"], leading=14, spaceAfter=8)
        italic = styles["Italic"]
        title_style = styles["Title"]

        story: List[Any] = []

        # Cover
        story.append(Spacer(1, 2.0 * inch))
        story.append(Paragraph(self._escape_pdf(title), title_style))
        story.append(Spacer(1, 0.3 * inch))
        subtitle = f"{metadata.get('genre','')}  {metadata.get('language','')}  {metadata.get('tone','')}"
        story.append(Paragraph(self._escape_pdf(subtitle.strip(" ")), italic))
        story.append(PageBreak())

        # Table of Contents (simple, eye-clean)
        story.append(Paragraph("Table of Contents", h1))
        story.append(Spacer(1, 0.15 * inch))
        for ch in outline.get("chapters", []):
            line = f"Chapter {ch['number']}: {ch['title']}"
            story.append(Paragraph(self._escape_pdf(line), body))
        story.append(PageBreak())

        # Chapters
        for ch in chapters:
            story.append(Paragraph(self._escape_pdf(f"Chapter {ch['number']}"), h2))
            story.append(Paragraph(self._escape_pdf(ch["title"]), h1))
            story.append(Spacer(1, 0.2 * inch))

            # Parse simple headings in content: "# " and "## "
            for block in self._split_blocks(ch["content"]):
                block = block.strip()
                if not block:
                    continue
                if block.startswith("# "):
                    story.append(Paragraph(self._escape_pdf(block[2:].strip()), h1))
                elif block.startswith("## "):
                    story.append(Paragraph(self._escape_pdf(block[3:].strip()), h2))
                else:
                    story.append(Paragraph(self._escape_pdf(block), body))

            story.append(PageBreak())

        def add_page_numbers(canvas, _doc):
            canvas.saveState()
            page_num = canvas.getPageNumber()
            canvas.setFont("Helvetica", 9)
            canvas.drawRightString(200 * mm, 15 * mm, str(page_num))
            canvas.restoreState()

        doc.build(story, onLaterPages=add_page_numbers)
        buf.seek(0)
        return buf.read()

    # -----------------------
    # DOCX export: editable book
    # -----------------------
    def _generate_book_docx(
        self,
        title: str,
        outline: Dict[str, Any],
        chapters: List[Dict[str, Any]],
        metadata: Dict[str, Any],
    ) -> bytes:
        d = Document()

        # Cover
        p = d.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(26)

        p2 = d.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub = f"{metadata.get('genre','')}  {metadata.get('language','')}  {metadata.get('tone','')}".strip(" ")
        r2 = p2.add_run(sub)
        r2.italic = True

        d.add_page_break()

        # TOC note (Word can auto-generate TOC from headings)
        d.add_heading("Table of Contents", level=1)
        d.add_paragraph(
            "NOTE: In Microsoft Word, you can insert an automatic Table of Contents from headings: "
            "References  Table of Contents. This document uses Heading styles for chapters/sections."
        )
        for ch in outline.get("chapters", []):
            d.add_paragraph(f"Chapter {ch['number']}: {ch['title']}")
        d.add_page_break()

        # Chapters with headings
        for ch in chapters:
            d.add_heading(f"Chapter {ch['number']}: {ch['title']}", level=1)

            for block in self._split_blocks(ch["content"]):
                block = block.strip()
                if not block:
                    continue
                if block.startswith("# "):
                    d.add_heading(block[2:].strip(), level=1)
                elif block.startswith("## "):
                    d.add_heading(block[3:].strip(), level=2)
                else:
                    d.add_paragraph(block)

            d.add_page_break()

        out = io.BytesIO()
        d.save(out)
        out.seek(0)
        return out.read()

    # -----------------------
    # Helpers (text parsing)
    # -----------------------
    def _split_blocks(self, text: str) -> List[str]:
        """
        Split text into paragraph-like blocks.
        We keep it simple and robust:
        - Split by blank lines
        - Preserve heading lines as their own blocks
        """
        if not text:
            return []
        # Normalize newlines
        t = text.replace("\r\n", "\n").replace("\r", "\n")
        blocks = re.split(r"\n\s*\n", t)
        return [b.strip() for b in blocks if b.strip()]

    def _escape_pdf(self, s: str) -> str:
        # Escape special chars for ReportLab Paragraph
        return (
            s.replace("&", "&amp;")
             .replace("<", "&lt;")
             .replace(">", "&gt;")
        )

    # -----------------------
    # Helpers (validation/convert)
    # -----------------------
    def _parse_json_input(self, v: Any, field: str) -> Any:
        if v is None:
            return None
        if isinstance(v, (dict, list)):
            return v
        if isinstance(v, str):
            v = v.strip()
            if not v:
                return None
            try:
                return json.loads(v)
            except Exception:
                raise ValueError(f"{field} must be valid JSON")
        raise ValueError(f"{field} must be a JSON object/array or string")

    def _clean_str(self, v: Any, max_len: int) -> str:
        s = "" if v is None else str(v).strip()
        if len(s) > max_len:
            s = s[:max_len]
        return s

    def _clean_required_str(self, v: Any, field: str, max_len: int) -> str:
        s = self._clean_str(v, max_len)
        if not s:
            raise ValueError(f"{field} is required")
        return s

    def _to_int_required(self, v: Any, field: str) -> int:
        if v is None or str(v).strip() == "":
            raise ValueError(f"{field} is required")
        try:
            return int(float(str(v).strip()))
        except Exception:
            raise ValueError(f"{field} must be a valid integer")

    def _to_int_optional(self, v: Any) -> Optional[int]:
        if v is None:
            return None
        s = str(v).strip()
        if not s:
            return None
        try:
            return int(float(s))
        except Exception:
            return None

    def _sanitize_filename(self, filename: str) -> str:
        invalid = '<>:"/\\|?*'
        for ch in invalid:
            filename = filename.replace(ch, "_")
        filename = filename.replace(" ", "_").strip(". ")
        return filename[:100] if filename else "book"

    def _ms_since(self, t_start: float) -> int:
        return int((time.perf_counter() - t_start) * 1000)

